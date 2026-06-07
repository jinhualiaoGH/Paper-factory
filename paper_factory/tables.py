from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_table(doc, table_data, table_number):
    title = table_data.get("title", f"Table {table_number}")
    headers = table_data.get("headers", [])
    rows = table_data.get("rows", [])

    if not headers:
        doc.add_paragraph(f"[Table {table_number} skipped: missing headers]")
        return

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"

    # Header row
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(header))
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)

    # Data rows
    for row in rows:
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            p = row_cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(value))
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)

    # Caption
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption.add_run(f"Table {table_number}. {title}")
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)

    doc.add_paragraph()