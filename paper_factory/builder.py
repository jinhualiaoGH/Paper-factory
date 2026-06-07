from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from paper_factory.pdf_export import export_pdf
from paper_factory.markdown_loader import markdown_to_sections
from paper_factory.bibliography import format_references
from paper_factory.discovery import discover_figures
from paper_factory.discovery import discover_csv_tables
from paper_factory.discovery import discover_references
from paper_factory.csv_loader import csv_to_table_data
from paper_factory.tables import add_table
from paper_factory.audit import audit_paper
import json
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def load_caption(base_name, caption_folder="captions"):
    caption_file = os.path.join(caption_folder, base_name + ".txt")

    if os.path.exists(caption_file):
        with open(caption_file, "r", encoding="utf-8") as f:
            return f.read().strip()

    return None

def load_config(path):
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


def set_default_font(doc, font_name="Times New Roman", font_size=12):
    styles = doc.styles
    normal_style = styles["Normal"]
    normal_style.font.name = font_name
    normal_style.font.size = Pt(font_size)

    for style_name in ["Heading 1", "Heading 2", "Title"]:
        if style_name in styles:
            styles[style_name].font.name = font_name


def set_margins(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)


def add_centered_paragraph(doc, text, font_size=12, bold=False, italic=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(font_size)
    run.bold = bold
    run.italic = italic
    return p


def add_figure(doc, figure_path, caption, number):
    if not os.path.exists(figure_path):
        doc.add_paragraph(f"[Missing Figure {number}: {figure_path}]")
        return

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(figure_path, width=Inches(5.5))

    add_centered_paragraph(
        doc,
        f"Figure {number}. {caption}",
        font_size=10,
        bold=False,
        italic=True
    )


def build_paper(config_path):

    cfg = load_config(config_path)
    if "markdown_file" in cfg:
        cfg["sections"] = markdown_to_sections(
            cfg["markdown_file"]
     )
        # Auto-discover CSV tables if none are defined
    if "tables" not in cfg:
        discovered_tables = discover_csv_tables("data")

        if discovered_tables:
            cfg["tables"] = discovered_tables

    # Auto-discover figures if none are defined
    if "figures" not in cfg:
        discovered_figures = discover_figures("figures")

        if discovered_figures:
            cfg["figures"] = discovered_figures
        # Auto-discover references if none are defined

    if "references" not in cfg:
        discovered_references = discover_references("references")

        if discovered_references:
            cfg["references"] = discovered_references

    issues = audit_paper(cfg)

    if issues:

        print()
        print("AUDIT REPORT")
        print("------------")

        for issue in issues:
            print(issue)

        print()

    else:
        print("Audit passed.")

    doc = Document()
    set_margins(doc)
    set_default_font(doc)

    add_centered_paragraph(doc, cfg["title"], font_size=16, bold=True)

    for author in cfg["authors"]:
        add_centered_paragraph(doc, author, font_size=12)

    add_centered_paragraph(doc, cfg["affiliation"], font_size=12)
    add_centered_paragraph(doc, cfg["email"], font_size=12)

    doc.add_paragraph()

    # -------------------------------------------------
    # Table of Contents page
    # -------------------------------------------------

    doc.add_page_break()

    p = doc.add_paragraph()
    run = p.add_run("Table of Contents")
    run.bold = True

    doc.add_paragraph(
        "(Automatic TOC will be added in a future version)"
    )

    doc.add_page_break()

    if "abstract" in cfg:
        doc.add_heading("Abstract", level=1)
        doc.add_paragraph(cfg["abstract"])

    for i, sec in enumerate(cfg["sections"], start=1):
        doc.add_heading(f"{i}. {sec['title']}", level=1)

        # Tables
    if "tables" in cfg:
        doc.add_heading("Tables", level=1)

    for i, table_data in enumerate(cfg["tables"], start=1):

        if "csv" in table_data:
            table_data = csv_to_table_data(
                table_data["csv"],
                table_data.get("title")
            )

        add_table(doc, table_data, i)

    if "figures" in cfg:
        doc.add_heading("Figures", level=1)
        for i, fig in enumerate(cfg["figures"], start=1):
            add_figure(doc, fig["file"], fig["caption"], i)
    
        # References
    if "references" in cfg:

        doc.add_page_break()

        doc.add_heading("References", level=1)
        style = cfg.get("bibliography_style", "IEEE")

        formatted_refs = format_references(
            cfg["references"],
            style
        )

        for ref in formatted_refs:
            doc.add_paragraph(ref)
 
    output_file = "output/paper.docx"
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0]

        p.alignment = 1  # center

        run = p.add_run("Page 1 ")

    #for section in doc.sections:
    #   footer = section.footer
    #   footer.is_linked_to_previous = False

    #    p = footer.paragraphs[0]
    #    p.text = "Page 1"
    #    p.alignment = 1

    doc.save(output_file)

    pdf_file = export_pdf(output_file)

    print()
    print("DOCX generated:")
    print(output_file)

    if pdf_file:
        print("PDF generated:")
        print(pdf_file)

 

def discover_references(folder="references"):
    references = []

    if not os.path.exists(folder):
        return references

    files = sorted(os.listdir(folder))

    for file in files:
        ext = os.path.splitext(file)[1].lower()

        if ext == ".txt":
            path = os.path.join(folder, file)

            with open(path, "r", encoding="utf-8") as f:
                ref = f.read().strip()

            if ref:
                references.append(ref)

    return references

def add_page_number(paragraph):
    paragraph.alignment = 1

    run = paragraph.add_run("Page ")

    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"

    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)